"""utility library used primarily for fixture generation in `conftest.py`.

contains no test to be run."""

import glob
import os
from os import path
import random
import re
import shutil
import zipfile

import docx
import jinja2
from spectrum import logger
from spectrum.config import COMMON

LOGGER = logger.logger(__name__)

def generate_article_id(msid):
    "given a regular 6-digit `msid`, generates a msid with a random prefix"
    msid = int(msid)

    # lsh@2022-08-3: offset increased from 100,000 to 1,000,000 as eLife msid approaches 100k
    offset = 1000000 # 10^6 for a 6 digit (or less) msid

    #     2^63 - 1 = 9223372036854775807 is the maximum id
    maximum_prefix = 9223372036854

    # special handling for the kitchen sink with it's large MSID
    kitchen_sink_id = 1234567890
    if msid == kitchen_sink_id:
        offset = 10000000000 # 10^10 for a 10 digit msid
        maximum_prefix = 92233720

    prefix = random.randrange(1, maximum_prefix + 1)
    return str((prefix * offset) + msid)

def article_zip(template_id, article_id=None, template_variables=None):
    if template_variables is None:
        template_variables = {}
    (template, kind) = _choose_template(template_id)
    if article_id is None:
        article_id = generate_article_id(template_id)
    generated_article_directory = '%s/elife-%s-%s-r1' % (COMMON['tmp'], article_id, kind)
    os.mkdir(generated_article_directory)
    generated_files = []
    for file in glob.glob(template + "/*"):
        generated_file = _generate(file, article_id, generated_article_directory, template_id, template_variables)
        generated_files.append(generated_file)
    zip_filename = generated_article_directory + '.zip'
    figure_names = []
    with zipfile.ZipFile(zip_filename, 'w') as zip_file:
        for generated_file in generated_files:
            zip_file.write(generated_file, path.basename(generated_file))
            match = re.match(r".*/elife-\d+-(.+).tif", generated_file)
            if match:
                figure_names.append(match.groups()[0])
    LOGGER.info("Generated %s with figures %s", zip_filename, figure_names, extra={'id': article_id})
    has_pdf = len(glob.glob(template + "/*.pdf")) >= 1
    return ArticleZip(article_id, zip_filename, generated_article_directory, revision=1, version=1, figure_names=figure_names, has_pdf=has_pdf)

def article_ejp_csv(source_csv, target_article_id, source_article_id=36157):
    generated_ejp_directory = '%s/poa-%s' % (COMMON['tmp'], target_article_id)
    if not path.exists(generated_ejp_directory):
        os.mkdir(generated_ejp_directory)
    generated_csv = path.join(generated_ejp_directory, path.basename(source_csv))
    with open(source_csv) as source:
        with open(generated_csv, 'w') as target:
            target.write(source.read().replace(
                str(source_article_id),
                target_article_id
            ))

    LOGGER.info("Generated EJP POA csv %s", generated_csv, extra={'id': target_article_id})

    return generated_csv

def article_ejp_zip(source_zip, target_article_id, source_article_id=36157):
    def _substitute_article_id(text):
        return re.sub(
            r"eLife\.%s\b" % source_article_id,
            "eLife.%s" % str(target_article_id),
            text
        )

    # 50142_1_supp_mat_highwire_zip_853595_pxvg3m.zip
    zip_prefix = random.randrange(1000000000, 9999999999 + 1)
    generated_ejp_zip_filename = path.join(
        COMMON['tmp'],
        "%s_1_supp_mat_highwire_zip_853595_abcdef.zip" % zip_prefix
    )

    with zipfile.ZipFile(source_zip, 'r') as source_zip_file:
        with zipfile.ZipFile(generated_ejp_zip_filename, 'w') as zip_file:
            for source_archived_filename in source_zip_file.namelist():
                with source_zip_file.open(source_archived_filename, 'r') as source_archived_file:
                    zip_file.writestr(source_archived_filename, _substitute_article_id(source_archived_file.read()))

    LOGGER.info("Generated EJP POA zip %s", generated_ejp_zip_filename, extra={'id': target_article_id})

    return generated_ejp_zip_filename

def digest_zip(template_id):
    standard_input = 'spectrum/templates/digests/DIGEST 99999.docx'
    article_id = generate_article_id(template_id)
    target_doc_filename = '%s/DIGEST %s.docx' % (COMMON['tmp'], article_id)

    word_document = docx.Document(standard_input)
    doi_paragraphs = [p for p in word_document.paragraphs if p.runs[0].text == 'MANUSCRIPT NUMBER\n']
    assert len(doi_paragraphs) == 1, "Wrong number of MANUSCRIPT NUMBER paragraphs: %s" % [p.text for p in word_document.paragraphs]
    doi_paragraphs[0].runs[1].text = '%s' % article_id
    word_document.save(target_doc_filename)

    LOGGER.info("Generated digest doc %s", target_doc_filename, extra={'id': article_id})

    generated_files = [target_doc_filename, 'spectrum/templates/digests/alligator.jpg']
    target_zip_filename = '%s/DIGEST %s.zip' % (COMMON['tmp'], article_id)
    with zipfile.ZipFile(target_zip_filename, 'w') as zip_file:
        for generated_file in generated_files:
            zip_file.write(generated_file, path.basename(generated_file))

    LOGGER.info("Generated digest zip %s", target_zip_filename, extra={'id': article_id})
    os.remove(target_doc_filename)

    return DigestZip(article_id, target_zip_filename)


def generate_article_title():
    return 'My spectrum article %s' % random.randrange(1, 1000000000000)


def clean():
    for entry in glob.glob('%s/elife*' % COMMON['tmp']):
        if path.isdir(entry):
            shutil.rmtree(entry)
            LOGGER.info("Deleted directory %s", entry)
        else:
            os.remove(entry)
            LOGGER.info("Deleted file %s", entry)

def all_stored_articles():
    """Returns all articles available as test inputs.

    However, excludes some blacklisted articles that we prefer to use with a specific test rather than the standard ingest-and-publish"""
    blacklist = ['19532', '06847', '22661']
    articles = []
    for template_directory in glob.glob('spectrum/templates/elife-*-*-*'):
        match = re.match(r".*/elife-(\d+)-.+", template_directory)
        assert match is not None, ("Cannot match an article id into %s" % template_directory)
        assert len(match.groups()) == 1
        article_id = match.groups()[0]
        if article_id in blacklist:
            continue
        articles.append(article_id)
    articles.sort()
    return articles

def _choose_template(template_id):
    templates_pattern = './spectrum/templates/elife-%s-*-*' % template_id
    templates_found = glob.glob(templates_pattern)
    assert len(templates_found) > 0, "No candidate templates found for: %s" % templates_pattern
    assert len(templates_found) == 1, "Found multiple candidate templates: %s" % templates_found
    chosen = templates_found[0]
    match = re.match(r'.*/elife-\d+-(vor|poa)-(r|v)\d+', chosen)
    assert match is not None, ("Bad name for template directory %s" % chosen)
    assert len(match.groups()) == 2
    kind = match.groups()[0] # vor or poa
    return (chosen, kind)


def _generate(filename, id, generated_article_directory, template_id, template_variables):
    filename_components = path.splitext(filename)
    generated_filename = path.basename(filename).replace(template_id, id)
    target = generated_article_directory + '/' + generated_filename
    assert len(filename_components) == 2
    extension = filename_components[1]
    if extension == '.jinja':
        with open(filename, 'r') as template_file:
            data = template_file.read()
        template = jinja2.Template(data)
        content = template.render(article={'id': id}, **template_variables)
        target = target.replace('.jinja', '')
        with open(target, 'w') as target_file:
            target_file.write(content)
    else:
        shutil.copy(filename, target)
    return target

class ArticleZip:
    def __init__(self, id, filename, directory, revision, version, figure_names=None, has_pdf=False):
        self._id = id
        self._filename = filename
        self._directory = directory
        self._revision = revision
        self._version = version
        self._figure_names = figure_names if figure_names else []
        self._has_pdf = has_pdf

    def id(self):
        return self._id

    def doi(self):
        return '10.7554/eLife.' + self._id

    def version(self):
        return self._version

    def filename(self):
        return self._filename

    def figure_names(self):
        return self._figure_names

    def has_pdf(self):
        return self._has_pdf

    def new_revision(self, version=None):
        if version:
            new_version = version
        else:
            new_version = self._version
        new_revision = self._revision + 1
        new_filename = re.sub(r'-(r|v)\d+.zip$', ('-r%s.zip' % new_revision), self._filename)
        shutil.copy(self._filename, new_filename)
        new_directory = re.sub(r'-(r|v)\d+$', ('-r%s' % new_revision), self._directory)
        shutil.copytree(self._directory, new_directory)
        return ArticleZip(self._id, new_filename, new_directory, new_revision, new_version, self._figure_names, self._has_pdf)

    def new_version(self, version):
        # what is changed is actually the "run"
        new_revision = self._revision + 1
        new_filename = re.sub(r'-(r|v)\d+.zip$', ('-v%s.zip' % version), self._filename)
        shutil.copy(self._filename, new_filename)
        new_directory = re.sub(r'-(r|v)\d+$', ('-v%s' % version), self._directory)
        shutil.copytree(self._directory, new_directory)
        return ArticleZip(self._id, new_filename, new_directory, new_revision, version, self._figure_names, self._has_pdf)

    def replace_in_text(self, replacements):
        """Beware: violates immutability, as it modifies the file in place for performance reasons"""
        LOGGER.info("Replacing %s in article", replacements, extra={'id': self._id})
        with zipfile.ZipFile(self._filename, 'w') as zip_file:
            for file in glob.glob(self._directory + "/*"):
                if file.endswith('.xml'):
                    with open(file) as xml:
                        contents = xml.read()
                    for search, replace in replacements.items():
                        contents = contents.replace(search, replace)
                    with open(file, 'w') as xml:
                        xml.write(contents)
                zip_file.write(file, path.basename(file))
        return self

    def clean(self):
        if os.path.exists(self._filename):
            os.remove(self._filename)
            LOGGER.info("Deleted file %s", self._filename)
        else:
            LOGGER.info("Not deleted file %s because it doesn't exist", self._filename)
        if os.path.exists(self._directory):
            shutil.rmtree(self._directory)
            LOGGER.info("Deleted directory %s", self._directory)
        else:
            LOGGER.info("Not deleted directory %s because it doesn't exist", self._directory)

def article_subjects(ids_to_subjects):
    maximum_suffix = 999999999
    suffix = random.randrange(1, maximum_suffix + 1)
    generated_file_name = '%s/article_subjects_%s.csv' % (COMMON['tmp'], suffix)
    with open(generated_file_name, 'w') as generated_file:
        generated_file.write('DOI,subj-group-type,subject\n')
        for id, subject in ids_to_subjects.items():
            generated_file.write('10.7554/eLife.%s,heading,%s\n' % (id, subject))
    return ArticleSubjects(generated_file_name)

class ArticleSubjects:
    def __init__(self, filename):
        self._filename = filename

    def filename(self):
        return self._filename

class DigestZip:
    def __init__(self, article_id, filename):
        self._article_id = article_id
        self._filename = filename

    def filename(self):
        return self._filename

    def article_id(self):
        return self._article_id
